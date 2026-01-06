"""
File Ingestion Module - Load documents into RAG for borgo-ai
Supports: PDF, TXT, Markdown, HTML, Word docs, CSV, JSON
"""
import os
import re
import json
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
import mimetypes

from config import KNOWLEDGE_DIR


@dataclass
class DocumentChunk:
    """A chunk of a document"""
    content: str
    metadata: Dict
    chunk_index: int
    total_chunks: int


@dataclass
class LoadedDocument:
    """A loaded document"""
    filename: str
    filepath: str
    content: str
    chunks: List[DocumentChunk]
    file_type: str
    size_bytes: int
    success: bool
    error: Optional[str] = None


class SemanticChunker:
    """
    Semantic chunking - splits by meaning, not just characters.
    
    Better than character chunking because:
    - Keeps paragraphs together
    - Respects section boundaries
    - Doesn't cut sentences in half
    - Preserves code blocks
    """
    
    def __init__(
        self,
        max_chunk_size: int = 1000,
        min_chunk_size: int = 100,
        overlap_sentences: int = 1
    ):
        self.max_chunk_size = max_chunk_size
        self.min_chunk_size = min_chunk_size
        self.overlap_sentences = overlap_sentences
    
    def chunk(self, text: str, metadata: Dict = None) -> List[DocumentChunk]:
        """Split text into semantic chunks"""
        metadata = metadata or {}
        
        # First, split into sections (by headers, double newlines, etc.)
        sections = self._split_into_sections(text)
        
        chunks = []
        current_chunk = ""
        
        for section in sections:
            # If section fits in current chunk, add it
            if len(current_chunk) + len(section) <= self.max_chunk_size:
                current_chunk += section + "\n\n"
            else:
                # Save current chunk if it has content
                if current_chunk.strip() and len(current_chunk) >= self.min_chunk_size:
                    chunks.append(current_chunk.strip())
                
                # If section itself is too large, split by sentences
                if len(section) > self.max_chunk_size:
                    sentence_chunks = self._split_large_section(section)
                    chunks.extend(sentence_chunks)
                    current_chunk = ""
                else:
                    current_chunk = section + "\n\n"
        
        # Don't forget the last chunk
        if current_chunk.strip() and len(current_chunk) >= self.min_chunk_size:
            chunks.append(current_chunk.strip())
        
        # Convert to DocumentChunk objects with overlap context
        doc_chunks = []
        for i, chunk_text in enumerate(chunks):
            # Add overlap from previous chunk (last sentence)
            if i > 0 and self.overlap_sentences > 0:
                prev_sentences = self._get_last_sentences(chunks[i-1], self.overlap_sentences)
                if prev_sentences:
                    chunk_text = f"[...] {prev_sentences}\n\n{chunk_text}"
            
            doc_chunks.append(DocumentChunk(
                content=chunk_text,
                metadata={**metadata, "chunk_index": i},
                chunk_index=i,
                total_chunks=len(chunks)
            ))
        
        return doc_chunks
    
    def _split_into_sections(self, text: str) -> List[str]:
        """Split text by semantic boundaries"""
        # Split by markdown headers
        sections = re.split(r'\n(?=#{1,6}\s)', text)
        
        # Further split by double newlines (paragraphs)
        all_sections = []
        for section in sections:
            if len(section) > self.max_chunk_size:
                # Split by paragraphs
                paragraphs = re.split(r'\n\s*\n', section)
                all_sections.extend(p.strip() for p in paragraphs if p.strip())
            else:
                if section.strip():
                    all_sections.append(section.strip())
        
        return all_sections
    
    def _split_large_section(self, text: str) -> List[str]:
        """Split a large section by sentences"""
        # Split by sentence endings
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current = ""
        
        for sentence in sentences:
            if len(current) + len(sentence) <= self.max_chunk_size:
                current += sentence + " "
            else:
                if current.strip():
                    chunks.append(current.strip())
                current = sentence + " "
        
        if current.strip():
            chunks.append(current.strip())
        
        return chunks
    
    def _get_last_sentences(self, text: str, n: int) -> str:
        """Get the last n sentences from text"""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        last_n = sentences[-n:] if len(sentences) >= n else sentences
        return " ".join(last_n)


class DocumentLoader:
    """Load various document types"""
    
    def __init__(self):
        self.chunker = SemanticChunker()
        self.supported_extensions = {
            '.txt': self._load_text,
            '.md': self._load_text,
            '.markdown': self._load_text,
            '.py': self._load_code,
            '.js': self._load_code,
            '.ts': self._load_code,
            '.java': self._load_code,
            '.cpp': self._load_code,
            '.c': self._load_code,
            '.h': self._load_code,
            '.json': self._load_json,
            '.csv': self._load_csv,
            '.html': self._load_html,
            '.htm': self._load_html,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
        }
    
    def load(self, filepath: str) -> LoadedDocument:
        """Load a document from file"""
        path = Path(filepath)
        
        if not path.exists():
            return LoadedDocument(
                filename=path.name,
                filepath=str(path),
                content="",
                chunks=[],
                file_type="unknown",
                size_bytes=0,
                success=False,
                error=f"File not found: {filepath}"
            )
        
        ext = path.suffix.lower()
        size = path.stat().st_size
        
        if ext not in self.supported_extensions:
            return LoadedDocument(
                filename=path.name,
                filepath=str(path),
                content="",
                chunks=[],
                file_type=ext,
                size_bytes=size,
                success=False,
                error=f"Unsupported file type: {ext}"
            )
        
        try:
            loader_func = self.supported_extensions[ext]
            content = loader_func(path)
            
            # Chunk the content
            metadata = {
                "source": path.name,
                "filepath": str(path),
                "file_type": ext
            }
            chunks = self.chunker.chunk(content, metadata)
            
            return LoadedDocument(
                filename=path.name,
                filepath=str(path),
                content=content,
                chunks=chunks,
                file_type=ext,
                size_bytes=size,
                success=True
            )
        
        except Exception as e:
            return LoadedDocument(
                filename=path.name,
                filepath=str(path),
                content="",
                chunks=[],
                file_type=ext,
                size_bytes=size,
                success=False,
                error=str(e)
            )
    
    def _load_text(self, path: Path) -> str:
        """Load plain text file"""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _load_code(self, path: Path) -> str:
        """Load code file with language annotation"""
        content = self._load_text(path)
        lang = path.suffix[1:]  # Remove the dot
        return f"```{lang}\n# File: {path.name}\n{content}\n```"
    
    def _load_json(self, path: Path) -> str:
        """Load JSON file"""
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return f"JSON content from {path.name}:\n```json\n{json.dumps(data, indent=2)}\n```"
    
    def _load_csv(self, path: Path) -> str:
        """Load CSV file"""
        try:
            import csv
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                rows = list(reader)
            
            if not rows:
                return "Empty CSV file"
            
            # Format as markdown table
            header = rows[0]
            content = f"CSV data from {path.name}:\n\n"
            content += "| " + " | ".join(header) + " |\n"
            content += "| " + " | ".join(["---"] * len(header)) + " |\n"
            
            for row in rows[1:50]:  # Limit to first 50 rows
                content += "| " + " | ".join(str(cell)[:50] for cell in row) + " |\n"
            
            if len(rows) > 51:
                content += f"\n... and {len(rows) - 51} more rows"
            
            return content
        except Exception as e:
            return f"Error loading CSV: {e}"
    
    def _load_html(self, path: Path) -> str:
        """Load HTML file and extract text"""
        try:
            from bs4 import BeautifulSoup
            
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # Remove scripts and styles
            for tag in soup(['script', 'style', 'nav', 'footer']):
                tag.decompose()
            
            title = soup.title.string if soup.title else path.name
            text = soup.get_text(separator='\n', strip=True)
            
            return f"# {title}\n\n{text}"
        except ImportError:
            # Fallback without BeautifulSoup
            content = self._load_text(path)
            # Basic HTML tag removal
            content = re.sub(r'<[^>]+>', '', content)
            return content
    
    def _load_pdf(self, path: Path) -> str:
        """Load PDF file"""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(str(path))
            text_parts = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- Page {i+1} ---\n{text}")
            
            return f"# PDF: {path.name}\n\n" + "\n\n".join(text_parts)
        
        except ImportError:
            return f"PDF support requires pypdf. Install with: pip install pypdf"
        except Exception as e:
            return f"Error loading PDF: {e}"
    
    def _load_docx(self, path: Path) -> str:
        """Load Word document"""
        try:
            from docx import Document
            
            doc = Document(str(path))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                        paragraphs.append(f"{'#' * level} {para.text}")
                    else:
                        paragraphs.append(para.text)
            
            return f"# Document: {path.name}\n\n" + "\n\n".join(paragraphs)
        
        except ImportError:
            return f"Word support requires python-docx. Install with: pip install python-docx"
        except Exception as e:
            return f"Error loading Word document: {e}"
    
    def load_directory(self, dirpath: str, recursive: bool = True) -> List[LoadedDocument]:
        """Load all supported documents from a directory"""
        path = Path(dirpath)
        documents = []
        
        if not path.exists():
            return documents
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                doc = self.load(str(file_path))
                documents.append(doc)
        
        return documents


def load_file(filepath: str) -> LoadedDocument:
    """Convenience function to load a file"""
    loader = DocumentLoader()
    return loader.load(filepath)


def load_directory(dirpath: str, recursive: bool = True) -> List[LoadedDocument]:
    """Convenience function to load a directory"""
    loader = DocumentLoader()
    return loader.load_directory(dirpath, recursive)
